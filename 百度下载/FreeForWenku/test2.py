import requests
import re
import argparse
import sys
import json
import os
from docx import Document
from docx.oxml.ns import qn

import re

# imgs to pdf
import glob
import fitz



# 图片转成PDF
def pic2pdf(doc_id,url):
    doc = fitz.open()
    for img in sorted(glob.glob(doc_id+"/*")):  # 读取图片，确保按文件名排序
        # 当文件出现错误的时候，自动跳过
        try:
            imgdoc = fitz.open(img)                 # 打开图片
            pdfbytes = imgdoc.convertToPDF()        # 使用图片创建单页的 PDF
            imgpdf = fitz.open("pdf", pdfbytes)
            doc.insertPDF(imgpdf)                   # 将当前页插入文档
        except:
            continue

    new_file_name = find_file_name(url)

    try:
        doc.save(new_file_name+".pdf")                   # 保存pdf文件
    except:
        doc.save(new_file_name + "-copy.pdf")  # 保存pdf文件
    doc.close()


# 删除指定的文件夹
def del_files(filepath):
    """
    删除某一目录下的所有文件或文件夹
    :param filepath: 路径
    :return:
    """
    del_list = os.listdir(filepath)
    for f in del_list:
        file_path = os.path.join(filepath, f)
        if os.path.isfile(file_path):
            os.remove(file_path)
        elif os.path.isdir(file_path):
            shutil.rmtree(file_path)

# 找的文档的名字
def find_file_name(url):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_12_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/62.0.3202.75 Safari/537.36'}

    # 获取页面
    source_html = None
    try:
        response = requests.get(url, headers=headers)
        source_html = response.content
    except Exception as e:
        print(e)

    # 解析源码
    try:
        content = source_html.decode('UTF-8')
    except:
        content = source_html.decode('gbk')
    # print(content)
    pattern = re.compile('<title>(.*?)</title>', re.S)
    title = re.findall(pattern, content)

    try:
        return title[0][:-7]
    except:
        return str(title)

#根据文件决定函数
y = 0
def DOC(url):
    doc_id = re.findall('view/(.*).html', url)[0]
    html = requests.get(url).text
    lists = re.findall('(https.*?0.json.*?)\\\\x22}', html)
    lenth = (len(lists) // 2)
    NewLists = lists[:lenth]
    for i in range(len(NewLists)):
        NewLists[i] = NewLists[i].replace('\\', '')
        txts = requests.get(NewLists[i]).text
        txtlists = re.findall('"c":"(.*?)".*?"y":(.*?),', txts)
        for i in range(0, len(txtlists)):
            global y
            print(txtlists[i][0].encode('utf-8').decode('unicode_escape', 'ignore'))
            if y != txtlists[i][1]:
                y = txtlists[i][1]
                n = '\n'
            else:
                n = ''

            filename = doc_id + '.txt'
            with open(filename, 'a', encoding='utf-8') as f:
                f.write(n + txtlists[i][0].encode('utf-8').decode('unicode_escape', 'ignore').replace('\\', ''))
            print("文档保存在" + filename)


    # # 创建一个doc文档
    # document = Document()
    # # 字体
    # document.styles['Normal'].font.name = u'宋体'
    # document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    #
    #
    # with open(filename,'r+',encoding='utf-8') as f:
    #     lines = f.readlines()
    #     for line in lines:
    #         # print(len(line))
    #         if len(line)<=3 or re.match('）',line):
    #             continue
    #         else:
    #             blank_line = 0
    #             # 第一行记录文件名
    #             if row==0:
    #                 new_file_name = line
    #                 row+=1
    #                 # 插入文件标题
    #                 document.add_heading(new_file_name, 0)
    #             else:
    #                 if re.match('A|B|C|D',line):
    #                     document.add_paragraph('      '+line)
    #                 else:
    #                     document.add_paragraph(line)
    #
    #     # 删除空白行
    #     for paragraphs in document.paragraphs:
    #         if paragraphs.text == "\n":
    #             paragraphs.clear()
    #
    #     new_file_name = find_file_name(url)
    #     try:
    #         document.save(new_file_name+'.docx')
    #     except:
    #         document.save(new_file_name+ '——copy.docx')
    #
    # f.close()
    # os.remove(filename)


def PPT(url):
    doc_id = re.findall('view/(.*).html',url)[0]
    url = "https://wenku.baidu.com/browse/getbcsurl?doc_id="+doc_id+"&pn=1&rn=99999&type=ppt"
    html = requests.get(url).text
    lists=re.findall('{"zoom":"(.*?)","page"',html)
    for i in range(0,len(lists)):
        lists[i] = lists[i].replace("\\",'')
    try:
        os.mkdir(doc_id)
    except:
        pass
    for i in range(0,len(lists)):
        img=requests.get(lists[i]).content
        with open(doc_id+'\img'+str(i)+'.jpg','wb') as m:
            m.write(img)
    # print("PPT图片保存在" + doc_id +"文件夹")

    pic2pdf(doc_id,url)
    os.removedirs(doc_id)


def TXT(url):
    doc_id = re.findall('view/(.*).html', url)[0]
    url = "https://wenku.baidu.com/api/doc/getdocinfo?callback=cb&doc_id="+doc_id
    html = requests.get(url).text
    md5 = re.findall('"md5sum":"(.*?)"',html)[0]
    pn = re.findall('"totalPageNum":"(.*?)"',html)[0]
    rsign = re.findall('"rsign":"(.*?)"',html)[0]
    NewUrl = 'https://wkretype.bdimg.com/retype/text/'+doc_id+'?rn='+pn+'&type=txt'+md5+'&rsign='+rsign
    txt = requests.get(NewUrl).text
    jsons = json.loads(txt)
    texts=re.findall("'c': '(.*?)',",str(jsons))
    print(texts)
    filename=doc_id+'.txt'
    with open(filename,'a',encoding='utf-8') as f:
        for i in range(0,len(texts)):
            texts[i] = texts[i].replace('\\r','\r')
            texts[i] = texts[i].replace('\\n','\n')

            f.write(texts[i])
    print("文档保存在" + filename)

    new_file_name = None
    # 创建一个doc文档
    document = Document()
    # 字体
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    row = 0
    with open(filename, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        for line in lines:
            # print(len(line))
            if len(line) <= 3 or re.match('）', line):
                continue
            else:
                blank_line = 0
                # 第一行记录文件名
                if row == 0:
                    new_file_name = line
                    row += 1
                    # 插入文件标题
                    document.add_heading(new_file_name, 0)
                else:
                    if re.match('A|B|C|D', line):
                        document.add_paragraph('      ' + line)
                    else:
                        document.add_paragraph(line)

        # 删除空白行
        for paragraphs in document.paragraphs:
            if paragraphs.text == "\n":
                paragraphs.clear()
        try:
            document.save(new_file_name[:-2] + '.docx')
        except:
            document.save(new_file_name[:-2] + '——copy.docx')

    f.close()
    os.remove(filename)

def PDF(url):
    doc_id = re.findall('view/(.*).html',url)[0]
    url = "https://wenku.baidu.com/browse/getbcsurl?doc_id="+doc_id+"&pn=1&rn=99999&type=pdf"
    html = requests.get(url).text
    lists=re.findall('{"zoom":"(.*?)","page"',html)
    for i in range(0,len(lists)):
        lists[i] = lists[i].replace("\\",'')
    try:
        os.mkdir(doc_id)
    except:
        pass
    for i in range(0,len(lists)):
        img=requests.get(lists[i]).content
        with open(doc_id+'\img'+str(i)+'.jpg','wb') as m:
            m.write(img)
    # print("FPD图片保存在" + doc_id + "文件夹")

    pic2pdf(doc_id,url)

    del_files(doc_id)

def main():
    # url = 'https://wenku.baidu.com/view/b4841a88a0116c175f0e4866.html'
    # DOC(url)

    url = 'https://wenku.baidu.com/view/ae5c0bf0f18583d048645948.html'
    DOC(url)


if __name__ == "__main__":
    main()

#DOC("https://wenku.baidu.com/view/b4841a88a0116c175f0e4866.html")